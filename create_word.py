from flask import Flask, send_file
from docx import Document
from bs4 import BeautifulSoup
import io

app = Flask(__name__)

@app.route('/download_document', methods=['GET'])
def download_document(response):
    # Cria um documento do Word
    doc = Document()
    
    # Simulando uma resposta HTML para demonstração
    response = "<h1>Título do Documento</h1><p>Este é um parágrafo de exemplo para o documento do Word.</p>"

    # Analisa o HTML
    soup = BeautifulSoup(response, 'html.parser')

    # Encontra e converte tags HTML em elementos do Word
    for tag in soup.find_all():
        if tag.name == 'p':
            # Adiciona parágrafo ao documento do Word
            doc.add_paragraph(tag.text)
        elif tag.name == 'h1':
            # Adiciona título de nível 1 ao documento do Word
            doc.add_heading(tag.text, level=1)
        # Adicione mais condições para outras tags HTML que deseja converter

    # Cria um buffer de memória para armazenar o arquivo DOCX
    file_buffer = io.BytesIO()
    doc.save(file_buffer)
    file_buffer.seek(0)

    # Envia o arquivo para o cliente para download
    return send_file(
        file_buffer,
        as_attachment=True,
        download_name="document.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )